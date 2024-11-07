from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import streamlit as st

# Função para definir a altura de uma linha
def set_row_height(row, height_cm):
    tr = row._tr
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # Multiplicando por 567 para converter cm para twips
    trHeight.set(qn('w:hRule'), 'exact')  # Define altura exata
    tr.append(trHeight)

# Função para aplicar largura exata a uma célula
def set_column_widths(table, col_widths):
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

# Função para aplicar formatação a um parágrafo
def apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0), space_after=Pt(0)):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after
    paragraph_format.line_spacing = 1  # Espaçamento simples
    paragraph.alignment = {'left': 0, 'center': 1, 'right': 2}.get(alignment, 1)

# Função para aplicar sombreamento a uma célula
def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Função para centralizar verticalmente o texto em uma célula
def set_vertical_alignment(cell, alignment='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)

# Função para adicionar bordas duplas a uma célula
def add_double_borders(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'double')  # Bordas duplas
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)



def create_custom_table(doc, itens_configurados, observacao):


    # Definir o número correto de linhas: cabeçalho (1) + itens (len(itens_configurados)) + linha do total (1)
    num_linhas = len(itens_configurados) + 2  # Cabeçalho + itens + linha de total


    # Criar a tabela com o número correto de linhas e 4 colunas
    table = doc.add_table(rows=num_linhas, cols=4)

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.left_indent = Cm(0)
    table.autofit = False  # Desativar autofit

    # Definir larguras fixas para as colunas
    col_widths = [Cm(1.25), Cm(1.25), Cm(8.88), Cm(3.09)]
    set_column_widths(table, col_widths)

    # Cabeçalho da tabela
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", f"Escopo do Fornecimento: {st.session_state['classificacao_estacao_subestacao']}", "Valor Total"]

    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Cor branca
        apply_paragraph_formatting(paragraph, alignment='center')

        # Aplicar formatação e cores no cabeçalho
        set_cell_shading(cell, '00543C')  # Cor verde escura
        add_double_borders(cell)
        set_vertical_alignment(cell, 'center')

    # Preenchendo os itens na tabela
    for idx, item in enumerate(itens_configurados, start=1):


        if idx >= len(table.rows) - 1:  # Garantir que o índice não ultrapasse o número de linhas

            break

        row = table.rows[idx]


        # Coluna "Item"
        row.cells[0].text = str(idx).zfill(2)  # Número do item
        apply_paragraph_formatting(row.cells[0].paragraphs[0], alignment='center')
        add_double_borders(row.cells[0])
        set_vertical_alignment(row.cells[0], 'center')
        set_vertical_alignment(row.cells[0], 'center')

        # Coluna "Qtde"
        quantidade = str(item.get('quantidade', 'N/A'))
        if quantidade.isdigit():
            quantidade = quantidade.zfill(2)  # Quantidade com dois algarismos
        row.cells[1].text = quantidade
        apply_paragraph_formatting(row.cells[1].paragraphs[0], alignment='center')
        add_double_borders(row.cells[1])
        set_vertical_alignment(row.cells[1], 'center') 

        # Coluna "Descrição"
        descricao = item.get('descricao', 'N/A')
        row.cells[2].text = descricao
        apply_paragraph_formatting(row.cells[2].paragraphs[0], alignment='left', space_before=Pt(2), space_after=Pt(2))
        add_double_borders(row.cells[2]) 

        # Aplicar bordas duplas a todas as células
        for row in table.rows:
            for cell in row.cells:
                add_double_borders(cell)

        # Mesclar todas as células da coluna "Valor Total" para o intervalo de linhas dos itens (exceto cabeçalho e linha de total)
        valor_total_cell = table.cell(1, 3)  # Primeira célula na coluna "Valor Total" para itens
        for row in table.rows[2:-1]:  # Itera sobre as linhas de itens (ignorando cabeçalho e linha de total)
            valor_total_cell = valor_total_cell.merge(row.cells[3])  # Mescla as células subsequentes

        # Definir o valor total na célula mesclada
        total = sum(item['valor_total'] for item in itens_configurados)
        valor_total_cell.text = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        apply_paragraph_formatting(valor_total_cell.paragraphs[0], alignment='center')
        set_vertical_alignment(valor_total_cell, 'center')
        add_double_borders(valor_total_cell)

        # Garantir que as células podem ajustar automaticamente a altura (não mexer diretamente na altura da linha)
        for cell in row.cells:
            # Configurar para que o texto se ajuste automaticamente
            cell.width = Cm(1.25)  # Definir largura para garantir que o conteúdo tenha espaço suficiente para o texto
            # Deixar as linhas com altura automática para "quebrar o texto"
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinhar texto à esquerda
            paragraph.word_wrap = True  # Ativar quebra automática de linha
            apply_paragraph_formatting(paragraph)  # Aplicar formatação geral

    # Linha do total (última linha da tabela)
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[3])  # Mescla as células até "Descrição"
    total_row.cells[3].text = f"Total (s/IPI): R$ {sum(item['valor_total'] for item in itens_configurados):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Definir altura da linha de total
    set_row_height(total_row, 1.1)
    set_row_height(header_row, 1.1)

    # Formatação para a linha de total
    for idx in [0, 3]:
        paragraph = total_row.cells[idx].paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='right', space_before=Pt(0))  # Espaçamento removido
        set_cell_shading(total_row.cells[idx], '00543C')  # Cor de fundo verde escura
        add_double_borders(total_row.cells[idx])
        set_vertical_alignment(total_row.cells[idx], 'center')
            # Adicionar bordas duplas a todas as células da tabela

    # Garantir que todas as células da coluna "Descrição" sejam alinhadas à esquerda
    for row in table.rows[1:-1]:  # Ignorar o cabeçalho e a última linha (linha do total)
        cell = row.cells[2]  # Coluna "Descrição"
        apply_paragraph_formatting(cell.paragraphs[0], alignment='left', space_before=Pt(2), space_after=Pt(2))

        return table



def set_table_left_indent(table, indent):
    tbl_pr = table._tbl.tblPr
    tbl_indent = tbl_pr.xpath("w:tblInd")
    if tbl_indent:
        tbl_indent[0].set(qn('w:w'), str(indent))
        tbl_indent[0].set(qn('w:type'), 'dxa')
    else:
        tbl_indent_element = OxmlElement('w:tblInd')
        tbl_indent_element.set(qn('w:w'), str(indent))
        tbl_indent_element.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_indent_element)

# Função para substituir texto no documento, incluindo tabelas e cabeçalhos
def substituir_texto_documento(doc, replacements):


    # Função para substituir texto em runs unidos de um parágrafo, preservando elementos não textuais
    def substituir_texto_paragrafo(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs if run.text)  # Combina todo o texto dos runs
        texto_modificado = full_text  # Copia inicial do texto completo

        for old_text, new_text in replacements.items():
            if old_text in texto_modificado:
                texto_modificado = texto_modificado.replace(old_text, new_text)  # Substitui no texto combinado

        # Distribuir o texto modificado de volta nos runs sem afetar outros elementos
        if texto_modificado != full_text:  # Apenas modifica se houve mudança
            texto_atualizado = texto_modificado  # Inicia com o texto modificado

            for run in paragraph.runs:
                if run.text:  # Somente atualiza o texto se houver conteúdo
                    if run.text in full_text:  # Confere se o run faz parte do texto original
                        run.text = texto_atualizado[:len(run.text)]  # Atualiza parte inicial do texto modificado
                        texto_atualizado = texto_atualizado[len(run.text):]  # Remove o texto atualizado

                # Preserva conteúdo não textual (e.g., imagens, quebras de linha/página)
                if run._r.xml.find('<a:blip') != -1:  # Verifica a presença de imagem (tag <a:blip>)
                    continue  # Não faz alterações no run com imagem

    # Substituir texto em todos os parágrafos do documento

    for paragraph in doc.paragraphs:
        substituir_texto_paragrafo(paragraph, replacements)

    # Substituir texto em todas as tabelas do documento

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_texto_paragrafo(paragraph, replacements)

    # Substituir texto no cabeçalho de todas as seções do documento

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            substituir_texto_paragrafo(paragraph, replacements)


def inserir_paragrafo_pagamento(doc, index):
    eventos_pagamento = st.session_state['eventos_pagamento']
    
    for produto, dados in eventos_pagamento.items():
        # Adicionar o nome do produto
        p = doc.add_paragraph()
        p.add_run(f"{produto}:\n").bold = True
        
        for evento in dados['eventos']:
            percentual = evento['percentual']
            dias = evento['dias']
            evento_texto = evento['evento']
            p.add_run(f" {percentual}%  - {dias} Dias da {evento_texto}\n")

    # Inserir o parágrafo no índice especificado
    doc.paragraphs[index]._element.addnext(p._element)

# Função para inserir as tabelas e realizar substituições de texto
def inserir_tabelas_word(doc, itens_configurados, observacao, replacements):
    # Realizar a substituição do texto usando o dicionário replacements

    # Procurar o parágrafo "Quadro de Preços" para inserir a tabela logo depois
    for i, paragraph in enumerate(doc.paragraphs):
        if "Quadro de Preços" in paragraph.text:
            # Inserir a tabela de Quadro de Preços logo após o parágrafo
            table = create_custom_table(doc, itens_configurados, observacao)
            doc.paragraphs[i+1]._element.addnext(table._element)
            break
    substituir_texto_documento(doc, replacements)

    # Retornar o documento atualizado
    return doc
    # Salvar o documento no caminho de saída
    doc.save(output_path)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Função para inserir as condições de pagamento
def inserir_eventos_pagamento(doc, eventos_pagamento):
    # Procurar o parágrafo "Condições de Pagamento"
    for i, paragraph in enumerate(doc.paragraphs):
        if "Condições de Pagamento" in paragraph.text:
            # Encontrado o parágrafo, inserir o novo parágrafo abaixo dele
            inserir_paragrafo_pagamento(doc, i + 1, eventos_pagamento)
            break

def inserir_paragrafo_pagamento(doc, index, eventos_pagamento):
    for produto, dados in eventos_pagamento.items():
        # Adicionar o nome do produto com formatação específica
        p = doc.add_paragraph()
        run = p.add_run(f"{produto}:\n")
        run.bold = True  # Tornar o nome do produto em negrito
        run.font.size = Pt(12)  # Tamanho da fonte 12
        run.font.name = 'Verdana'  # Fonte Verdana
        run.font.color.rgb = RGBColor(0, 84, 60)  # Cor #00543C
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinhar à esquerda
        
        # Adicionar as condições de pagamento para cada evento
        for evento in dados['eventos']:
            percentual = int(evento['percentual'])
            dias = evento['dias']
            evento_texto = evento['evento']
            
            # Verificar se dias é 0 e substituir por "com a"
            if dias == 0:
                dias_texto = "com a"  # Substituir 0 por "com a" sem a palavra "Dias" ou "da"
            else:
                dias_texto = f"{dias} Dias da"  # Caso contrário, manter o valor com "Dias" e "da"
            
            # Criar um parágrafo para cada condição, aplicando o estilo de lista
            p_evento = doc.add_paragraph(style='Bullet')  # Estilo de marcador
            p_evento.add_run(f"{percentual}% - {dias_texto} {evento_texto}")
            p_evento.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinhar à esquerda

            # Inserir o parágrafo do evento
            doc.paragraphs[index]._element.addnext(p_evento._element)
        
        # Inserir o parágrafo com o produto e seus eventos na posição especificada
        doc.paragraphs[index]._element.addnext(p._element)

        # Adicionar mais parágrafos vazios para aumentar o espaço entre os blocos de produtos
        doc.add_paragraph()  # Parágrafo vazio adicional
        doc.add_paragraph() 

# Função para inserir a seção de impostos no documento
def inserir_impostos(doc, classificacao_estacao_subestacao):
    # Definir NCM e lista de impostos com base na classificação
    if classificacao_estacao_subestacao == "Estação Fotovoltaica":
        ncm_texto = "NCM: 8501.72.90 – Sistema Gerador Fotovoltaico"
        impostos = [
            "PIS: 1,65% inclusos nos preços;",
            "COFINS: 7,60% inclusos nos preços;",
            "ICMS: isento;",
            "IPI: isento."
        ]
    elif classificacao_estacao_subestacao == "Subestação Unitária":
        ncm_texto = "NCM: 8537.20.90 – Subestação Unitária"
        impostos = [
                    "PIS: 1,65% inclusos nos preços;"
                                "COFINS: 7,6% inclusos nos preços;",
                                                f"ICMS: {st.session_state['icms']+st.session_state['difal']+st.session_state['f_pobreza']}% inclusos nos preços;",

            "IPI: 0,00% a incluir nos preços.",


        
        ]
    else:
        return  # Saia da função se a classificação for inválida

    # Procurar o parágrafo "Impostos" e inserir o NCM e impostos abaixo dele
    for i, paragraph in enumerate(doc.paragraphs):
        if "Impostos" in paragraph.text:
            # Encontrado o parágrafo, chamar a função de inserção de parágrafos formatados
            inserir_paragrafo_impostos(doc, i , ncm_texto, impostos)
            break

# Função para formatar e inserir o NCM e a lista de impostos no documento
def inserir_paragrafo_impostos(doc, index, ncm_texto, impostos):
    # Adicionar o NCM com formatação específica
    p_ncm = doc.add_paragraph()
    run_ncm = p_ncm.add_run(ncm_texto)
    run_ncm.bold = True
    run_ncm.font.size = Pt(12)
    run_ncm.font.name = 'Verdana'
    run_ncm.font.color.rgb = RGBColor(0, 84, 60)
    p_ncm.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Adicionar o parágrafo formatado do NCM abaixo do índice especificado
    doc.paragraphs[index]._element.addnext(p_ncm._element)



    # Adicionar cada imposto na ordem correta como um item de lista bullet
    for imposto in impostos:
        p_imposto = doc.add_paragraph(style='Bullet')  # Usar estilo de lista de marcadores
        p_imposto_run = p_imposto.add_run(imposto)
        p_imposto_run.font.size = Pt(11)
        p_imposto_run.font.name = 'Calibri Light'  # Fonte "Calibri Light"
        p_imposto.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Inserir cada parágrafo do imposto logo após o último inserido
        doc.paragraphs[index + 1]._element.addnext(p_imposto._element)


